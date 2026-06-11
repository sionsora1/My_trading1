"""
Word/WPS文档身份证号提取工具

功能：
  - 读取 .docx / .doc / .wps 文档，用正则提取所有中国身份证号
  - 去重后输出到 Excel，每个号码占一行

用法：
  python extract_idcard.py <文件或目录路径> [--output 输出文件名]

依赖：
  pip install python-docx openpyxl
"""

import re
import sys
import os
from pathlib import Path


# ============================================================
# 身份证号正则
# ============================================================

# 18位：6位地区 + 8位生日(YYYYMMDD) + 3位顺序码 + 1位校验码(0-9/X)
ID18 = r'(?<!\d)[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx](?!\d)'

# 15位（旧版）：6位地区 + 6位生日(YYMMDD) + 3位顺序码
ID15 = r'(?<!\d)[1-9]\d{5}\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}(?!\d)'

ID_PATTERN = re.compile(f'{ID18}|{ID15}')


# ============================================================
# 文档读取
# ============================================================

def _extract_text_from_binary(filepath: str, verbose: bool = False) -> str:
    """
    从二进制/旧格式文件中暴力提取可读文本（回退方案）
    依次尝试 gbk → gb2312 → utf-16le → utf-16be → utf-8 → latin-1
    """
    try:
        with open(filepath, 'rb') as f:
            raw = f.read()
    except Exception:
        return ''

    if verbose:
        print(f"    文件大小: {len(raw)} bytes")

    # 方案A：尝试多种编码解码全文
    best = ''
    for enc in ['gbk', 'gb2312', 'utf-16-le', 'utf-16-be', 'utf-8', 'latin-1']:
        try:
            text = raw.decode(enc, errors='ignore')
            cn_count = sum(1 for c in text if '一' <= c <= '鿿')
            if cn_count > sum(1 for c in best if '一' <= c <= '鿿'):
                best = text
        except Exception:
            continue

    # 方案B：latin-1 兜底保留所有字节值（确保数字 0-9 不被破坏）
    latin1_text = raw.decode('latin-1', errors='ignore')
    if verbose:
        cn_best = sum(1 for c in best if '一' <= c <= '鿿')
        digit_runs = sum(1 for c in latin1_text if c.isdigit())
        print(f"    最佳编码: {len(best)} 字符, 中文 {cn_best} 个")
        print(f"    latin-1 回退: {len(latin1_text)} 字符, 含数字 {digit_runs} 个")

    # 合并两种提取结果
    combined = best
    if latin1_text != best:
        combined = best + '\n' + latin1_text

    if verbose and len(combined) < 1000:
        preview = combined[:300].replace('\n', ' ').replace('\r', '')
        print(f"    文本预览: {preview}")

    return combined


def extract_from_document(filepath: str, verbose: bool = False) -> list[str]:
    """
    从文档文件提取身份证号
    支持 .docx / .doc / .wps（新版 OOXML 格式或旧版二进制格式）
    """
    from docx import Document

    results = []

    # 方案1：用 python-docx 打开（支持 .docx 和新版 .wps）
    try:
        doc = Document(filepath)
        para_count = len(doc.paragraphs)
        for para in doc.paragraphs:
            results.extend(ID_PATTERN.findall(para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    results.extend(ID_PATTERN.findall(cell.text))
        if verbose:
            print(f"    python-docx 打开成功: {para_count} 段落, {len(doc.tables)} 表格, 找到 {len(results)} 个号码")
        # 只有真正找到内容才返回，否则继续尝试二进制回退
        if results:
            return results
    except Exception as e:
        if verbose:
            print(f"    python-docx 打不开: {e}")

    # 方案2：二进制回退（旧版 .doc / .wps 格式）
    if verbose:
        print(f"    尝试二进制回退...")
    text = _extract_text_from_binary(filepath, verbose=verbose)
    if text:
        results.extend(ID_PATTERN.findall(text))

    return results


def extract_from_txt(filepath: str) -> list[str]:
    """从纯文本文件提取身份证号"""
    results = []
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            text = f.read()
    except UnicodeDecodeError:
        try:
            with open(filepath, 'r', encoding='gbk') as f:
                text = f.read()
        except Exception as e:
            print(f"  [跳过] {filepath} — 编码错误: {e}")
            return results
    except Exception as e:
        print(f"  [跳过] {filepath} — 无法读取: {e}")
        return results

    results.extend(ID_PATTERN.findall(text))
    return results


# ============================================================
# 文件扫描
# ============================================================

def collect_files(path: str) -> list[str]:
    """收集所有支持的文件（含 .wps）"""
    p = Path(path)
    if not p.exists():
        print(f"[错误] 路径不存在: {path}")
        sys.exit(1)

    if p.is_file():
        return [str(p)]

    files = []
    for ext in ['*.docx', '*.doc', '*.wps', '*.txt']:
        files.extend(str(f) for f in p.rglob(ext))
    return sorted(files)


# ============================================================
# 主流程
# ============================================================

def main():
    import argparse

    parser = argparse.ArgumentParser(description='从Word/WPS文档提取身份证号，输出到Excel')
    parser.add_argument('path', help='文件路径或目录路径')
    parser.add_argument('--output', '-o', default='idcard_output.xlsx', help='输出Excel文件名 (默认: idcard_output.xlsx)')
    parser.add_argument('--verbose', '-v', action='store_true', help='显示详细提取过程（调试用）')
    args = parser.parse_args()

    # 1. 收集文件
    files = collect_files(args.path)
    if not files:
        print(f"未找到支持的文档文件 (.docx/.doc/.wps/.txt)")
        return

    print(f"找到 {len(files)} 个文件")

    # 2. 逐文件提取
    all_ids = []
    for f in files:
        print(f"  处理: {f}")
        ext = Path(f).suffix.lower()
        if ext in ('.docx', '.doc', '.wps'):
            ids = extract_from_document(f, verbose=args.verbose)
        else:
            ids = extract_from_txt(f)

        if ids:
            print(f"    → 找到 {len(ids)} 个身份证号")
        elif args.verbose:
            print(f"    → 未找到身份证号")
        all_ids.extend(ids)

    # 3. 去重（保持出现顺序）
    seen = set()
    unique_ids = []
    for id_num in all_ids:
        if id_num not in seen:
            seen.add(id_num)
            unique_ids.append(id_num)

    # 4. 输出
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "身份证号"
    ws['A1'] = '序号'
    ws['B1'] = '身份证号'
    ws['C1'] = '长度'
    ws['D1'] = '性别'

    for i, id_num in enumerate(unique_ids, start=2):
        ws.cell(row=i, column=1, value=i - 1)
        ws.cell(row=i, column=2, value=f"'{id_num}")  # 文本格式防止科学计数法
        ws.cell(row=i, column=3, value=len(id_num))
        # 性别：18位倒数第2位奇数=男，15位最后1位
        if len(id_num) == 18:
            gender = '男' if int(id_num[-2]) % 2 == 1 else '女'
        elif len(id_num) == 15:
            gender = '男' if int(id_num[-1]) % 2 == 1 else '女'
        else:
            gender = ''
        ws.cell(row=i, column=4, value=gender)

    # 列宽
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 22
    ws.column_dimensions['C'].width = 8
    ws.column_dimensions['D'].width = 8

    wb.save(args.output)

    print(f"\n{'='*50}")
    print(f"  提取完成！")
    print(f"  文件数: {len(files)}")
    print(f"  提取总数: {len(all_ids)}")
    print(f"  去重后: {len(unique_ids)}")
    print(f"  已保存: {os.path.abspath(args.output)}")
    print(f"{'='*50}")


if __name__ == '__main__':
    main()
