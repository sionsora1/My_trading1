"""
Word/WPS文档身份证号提取工具（可视化界面）

功能：
  - 图形界面选择 Word/WPS 文件或目录
  - 提取 18位/15位 身份证号，去重后导出 Excel
  - 显示提取进度和结果统计
  - 支持 .docx / .doc / .wps / .txt

用法：
  python extract_idcard_gui.py
"""

import re
import os
import sys
import threading
from pathlib import Path

# ============================================================
# 身份证号正则
# ============================================================
ID18 = r'(?<!\d)[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx](?!\d)'
ID15 = r'(?<!\d)[1-9]\d{5}\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}(?!\d)'
ID_PATTERN = re.compile(f'{ID18}|{ID15}')


# ============================================================
# 提取逻辑
# ============================================================
def _extract_text_from_binary(filepath: str) -> str:
    """从二进制/旧格式文件中暴力提取可读文本
    依次尝试 gbk / gb2312 / utf-16le / utf-16be / utf-8 / latin-1
    latin-1 作为兜底保留所有字节，确保数字不会被破坏"""
    try:
        with open(filepath, 'rb') as f:
            raw = f.read()
    except Exception:
        return ''

    best = ''
    for enc in ['gbk', 'gb2312', 'utf-16-le', 'utf-16-be', 'utf-8', 'latin-1']:
        try:
            text = raw.decode(enc, errors='ignore')
            cn = sum(1 for c in text if '一' <= c <= '鿿')
            if cn > sum(1 for c in best if '一' <= c <= '鿿'):
                best = text
        except Exception:
            continue

    latin1_text = raw.decode('latin-1', errors='ignore')
    if latin1_text != best and len(latin1_text) > 0:
        best = best + '\n' + latin1_text

    return best


def extract_from_document(filepath: str) -> list[str]:
    """从文档提取身份证号，支持 .docx / .doc / .wps"""
    from docx import Document

    results = []

    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            results.extend(ID_PATTERN.findall(para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    results.extend(ID_PATTERN.findall(cell.text))
        if results:
            return results
    except Exception:
        pass

    text = _extract_text_from_binary(filepath)
    if text:
        results.extend(ID_PATTERN.findall(text))

    return results


def extract_from_txt(filepath: str) -> list[str]:
    """从 .txt 提取"""
    results = []
    for enc in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
        try:
            with open(filepath, 'r', encoding=enc) as f:
                text = f.read()
            results.extend(ID_PATTERN.findall(text))
            break
        except (UnicodeDecodeError, Exception):
            continue
    return results


def collect_files(path: str) -> list[str]:
    """递归收集支持的文件（含 .wps）"""
    p = Path(path)
    if p.is_file():
        return [str(p)]
    files = []
    for ext in ['*.docx', '*.doc', '*.wps', '*.txt']:
        files.extend(str(f) for f in p.rglob(ext))
    return sorted(files)


# ============================================================
# GUI（tkinter）
# ============================================================
import tkinter as tk
from tkinter import ttk, filedialog, messagebox


class IDCardExtractorApp:
    def __init__(self, root):
        self.root = root
        self.root.title('身份证号提取工具')
        self.root.geometry('620x420')
        self.root.resizable(False, False)

        style = ttk.Style()
        style.theme_use('clam')

        self._build_ui()

    def _build_ui(self):
        title = ttk.Label(
            self.root,
            text='Word / WPS 文档身份证号提取',
            font=('Microsoft YaHei', 14, 'bold')
        )
        title.pack(pady=(20, 10))

        # 输入
        input_frame = ttk.LabelFrame(self.root, text='输入设置', padding=10)
        input_frame.pack(fill='x', padx=20, pady=(0, 10))

        row1 = ttk.Frame(input_frame)
        row1.pack(fill='x', pady=(0, 8))
        ttk.Label(row1, text='文件/目录:', width=10).pack(side='left')
        self.path_var = tk.StringVar()
        ttk.Entry(row1, textvariable=self.path_var, width=42).pack(side='left', padx=(0, 5))
        ttk.Button(row1, text='选文件', command=self._select_file, width=8).pack(side='left', padx=(0, 3))
        ttk.Button(row1, text='选目录', command=self._select_folder, width=8).pack(side='left')

        # 输出
        output_frame = ttk.LabelFrame(self.root, text='输出设置', padding=10)
        output_frame.pack(fill='x', padx=20, pady=(0, 10))

        row2 = ttk.Frame(output_frame)
        row2.pack(fill='x', pady=(0, 5))
        ttk.Label(row2, text='输出文件:', width=10).pack(side='left')
        self.output_var = tk.StringVar(value='idcard_output.xlsx')
        ttk.Entry(row2, textvariable=self.output_var, width=42).pack(side='left', padx=(0, 5))
        ttk.Button(row2, text='另存为...', command=self._select_output, width=8).pack(side='left')

        # 执行
        btn_frame = ttk.Frame(self.root)
        btn_frame.pack(pady=10)
        self.run_btn = ttk.Button(btn_frame, text='开始提取', command=self._start_extract, width=20)
        self.run_btn.pack()

        # 进度
        self.progress = ttk.Progressbar(self.root, mode='indeterminate', length=500)
        self.progress.pack(pady=(5, 5))

        # 状态
        self.status_var = tk.StringVar(value='就绪 - 请选择文件或目录')
        ttk.Label(self.root, textvariable=self.status_var, foreground='gray').pack()

        # 结果
        self.result_text = tk.Text(self.root, height=6, width=68, state='disabled', font=('Consolas', 9))
        self.result_text.pack(pady=(8, 10))

    def _select_file(self):
        path = filedialog.askopenfilename(
            title='选择 Word/WPS 文档',
            filetypes=[
                ('Word/WPS文档', '*.docx *.doc *.wps'),
                ('文本文件', '*.txt'),
                ('所有文件', '*.*'),
            ]
        )
        if path:
            self.path_var.set(path)
            base = Path(path).stem
            self.output_var.set(f'{base}_身份证号.xlsx')

    def _select_folder(self):
        path = filedialog.askdirectory(title='选择目录')
        if path:
            self.path_var.set(path)

    def _select_output(self):
        path = filedialog.asksaveasfilename(
            title='保存输出文件',
            defaultextension='.xlsx',
            filetypes=[('Excel文件', '*.xlsx')],
            initialfile=self.output_var.get(),
        )
        if path:
            self.output_var.set(path)

    def _start_extract(self):
        src = self.path_var.get().strip()
        if not src:
            messagebox.showwarning('提示', '请先选择 Word 文件或目录')
            return
        if not os.path.exists(src):
            messagebox.showerror('错误', f'路径不存在:\n{src}')
            return

        out = self.output_var.get().strip()
        if not out:
            messagebox.showwarning('提示', '请指定输出文件名')
            return

        self.run_btn.config(state='disabled')
        self.progress.start()
        self._set_result('正在提取...')
        self.status_var.set('处理中，请稍候...')

        thread = threading.Thread(target=self._do_extract, args=(src, out), daemon=True)
        thread.start()

    def _do_extract(self, src: str, out: str):
        try:
            files = collect_files(src)
            if not files:
                self._done(lambda: messagebox.showinfo('提示', '未找到 .docx / .doc / .wps / .txt 文件'))
                return

            all_ids = []
            for f in files:
                ext = Path(f).suffix.lower()
                if ext in ('.docx', '.doc', '.wps'):
                    ids = extract_from_document(f)
                else:
                    ids = extract_from_txt(f)
                all_ids.extend(ids)

            seen = set()
            unique = []
            for i in all_ids:
                if i not in seen:
                    seen.add(i)
                    unique.append(i)

            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = '身份证号'
            ws.append(['序号', '身份证号', '长度', '性别'])

            for idx, id_num in enumerate(unique, 1):
                length = len(id_num)
                if length == 18:
                    gender = '男' if int(id_num[-2]) % 2 == 1 else '女'
                elif length == 15:
                    gender = '男' if int(id_num[-1]) % 2 == 1 else '女'
                else:
                    gender = ''
                ws.append([idx, f"'{id_num}", length, gender])

            ws.column_dimensions['A'].width = 8
            ws.column_dimensions['B'].width = 22
            ws.column_dimensions['C'].width = 8
            ws.column_dimensions['D'].width = 8
            wb.save(out)

            result = (
                f'文件数:    {len(files)}\n'
                f'提取总数:  {len(all_ids)}\n'
                f'去重后:    {len(unique)}\n'
                f'输出文件:  {os.path.abspath(out)}\n'
                f'-' * 30
            )
            self._done(lambda r=result: self._finish(r))

        except Exception as e:
            self._done(lambda e=e: self._error(str(e)))

    def _set_result(self, text):
        self.root.after(0, lambda: (
            self.result_text.configure(state='normal'),
            self.result_text.delete('1.0', 'end'),
            self.result_text.insert('1.0', text),
            self.result_text.configure(state='disabled')
        ))

    def _finish(self, text):
        self.progress.stop()
        self.run_btn.config(state='normal')
        self.status_var.set('提取完成!')
        self._set_result(text)
        messagebox.showinfo('完成', '身份证号提取完成!')

    def _error(self, msg):
        self.progress.stop()
        self.run_btn.config(state='normal')
        self.status_var.set('提取失败')
        self._set_result(f'错误: {msg}')
        messagebox.showerror('错误', msg)

    def _done(self, callback):
        self.root.after(0, callback)


def main():
    root = tk.Tk()
    app = IDCardExtractorApp(root)
    root.mainloop()


if __name__ == '__main__':
    main()
